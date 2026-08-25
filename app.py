import streamlit as st
import openai
from io import BytesIO
from docx import Document
from datetime import datetime, timezone
import uuid
import json
import base64


# ============================================================
# PAGE CONFIG
# ============================================================

st.set_page_config(
    page_title="Primer Paso AI | Immigration Case Intelligence",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded",
)


# ============================================================
# CONSTANTS
# ============================================================

MODEL = "gpt-4o-mini"
VISION_MODEL = "gpt-4o"

TARGET_LANGUAGES = [
    "English",
    "Spanish",
    "French",
    "Haitian Creole",
    "Portuguese",
    "Mandarin",
]

PROTECTED_GROUNDS = [
    "Race",
    "Religion",
    "Nationality",
    "Membership in a particular social group (PSG)",
    "Political opinion",
]


# ============================================================
# CASE MEMORY
# ============================================================

def create_case():
    return {
        "case_id": str(uuid.uuid4()),
        "case_type": "Humanitarian Immigration",
        "applicant_name": "",
        "country_of_origin": "",
        "protected_ground": "",
        "sources": [],
        "facts": [],
        "timeline": [],
        "evidence": [],
        "authorities": [],
        "vulnerabilities": [],
        "outputs": [],
        "audit_log": [],
        "approvals": [],
    }


def initialize_session():
    if "case" not in st.session_state:
        st.session_state.case = create_case()

    if "last_output" not in st.session_state:
        st.session_state.last_output = ""

    if "workflow_result" not in st.session_state:
        st.session_state.workflow_result = None


initialize_session()


def get_case():
    return st.session_state.case


# ============================================================
# AUDIT LOG
# ============================================================

def audit_event(action, actor="AI_AGENT", details=None):
    event = {
        "event_id": str(uuid.uuid4()),
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "action": action,
        "actor": actor,
        "details": details or {},
    }

    st.session_state.case["audit_log"].append(event)

    return event


# ============================================================
# CASE HELPERS
# ============================================================

def add_source(
    source_type,
    name,
    text="",
    metadata=None,
):
    source = {
        "source_id": str(uuid.uuid4()),
        "source_type": source_type,
        "name": name,
        "text": text,
        "metadata": metadata or {},
        "created_at": datetime.now(timezone.utc).isoformat(),
    }

    st.session_state.case["sources"].append(source)

    audit_event(
        "SOURCE_ADDED",
        details={
            "source_id": source["source_id"],
            "source_type": source_type,
            "name": name,
        },
    )

    return source


def add_fact(
    statement,
    source_ids,
    date_value=None,
    confidence="medium",
    verified=False,
):
    fact = {
        "fact_id": str(uuid.uuid4()),
        "statement": statement,
        "date": date_value,
        "confidence": confidence,
        "source_ids": source_ids,
        "attorney_verified": verified,
    }

    st.session_state.case["facts"].append(fact)

    return fact


def add_timeline_event(
    date_value,
    description,
    fact_ids,
    source_ids,
    confidence="medium",
):
    event = {
        "event_id": str(uuid.uuid4()),
        "date": date_value,
        "description": description,
        "fact_ids": fact_ids,
        "source_ids": source_ids,
        "confidence": confidence,
    }

    st.session_state.case["timeline"].append(event)

    return event


def add_vulnerability(
    issue,
    severity,
    explanation,
    fact_ids=None,
    recommended_action=None,
):
    vulnerability = {
        "vulnerability_id": str(uuid.uuid4()),
        "issue": issue,
        "severity": severity,
        "explanation": explanation,
        "fact_ids": fact_ids or [],
        "recommended_action": recommended_action,
    }

    st.session_state.case["vulnerabilities"].append(vulnerability)

    return vulnerability


# ============================================================
# OPENAI CLIENT
# ============================================================

def get_openai_client():
    api_key = st.secrets.get("OPENAI_API_KEY")

    if not api_key:
        return None

    return openai.OpenAI(api_key=api_key)


client = get_openai_client()


# ============================================================
# AI HELPERS
# ============================================================

def ask_markdown(
    system_prompt,
    user_prompt,
    model=MODEL,
):
    if client is None:
        raise RuntimeError(
            "OPENAI_API_KEY is not configured in Streamlit Secrets."
        )

    response = client.chat.completions.create(
        model=model,
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": system_prompt,
            },
            {
                "role": "user",
                "content": user_prompt,
            },
        ],
    )

    return response.choices[0].message.content


def ask_json(
    system_prompt,
    user_prompt,
    model=MODEL,
):
    if client is None:
        raise RuntimeError(
            "OPENAI_API_KEY is not configured in Streamlit Secrets."
        )

    response = client.chat.completions.create(
        model=model,
        temperature=0,
        response_format={
            "type": "json_object"
        },
        messages=[
            {
                "role": "system",
                "content": system_prompt,
            },
            {
                "role": "user",
                "content": user_prompt,
            },
        ],
    )

    raw = response.choices[0].message.content

    try:
        return json.loads(raw)
    except json.JSONDecodeError as exc:
        raise ValueError(
            "The AI returned invalid structured data. "
            "No case data was saved."
        ) from exc


# ============================================================
# DOCUMENT GENERATION
# ============================================================

def create_docx(
    title,
    body,
):
    document = Document()

    document.add_heading(
        title,
        level=1,
    )

    document.add_paragraph(
        body
    )

    output = BytesIO()

    document.save(output)

    output.seek(0)

    return output


# ============================================================
# IMAGE OCR / TRANSLATION
# ============================================================

def process_image(
    image_bytes,
    mime_type,
    target_language,
):
    encoded = base64.b64encode(
        image_bytes
    ).decode("utf-8")

    response = client.chat.completions.create(
        model=VISION_MODEL,
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a source-faithful legal document "
                    "OCR and translation assistant.\n\n"
                    "Perform these steps separately:\n"
                    "1. Extract the visible text.\n"
                    "2. Translate the extracted text.\n"
                    "3. Identify possible legal relevance.\n\n"
                    "Do not invent missing text.\n"
                    "If text is unreadable, say so.\n"
                    "Do not make legal conclusions."
                ),
            },
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            f"Target translation language: "
                            f"{target_language}"
                        ),
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": (
                                f"data:{mime_type};"
                                f"base64,{encoded}"
                            )
                        },
                    },
                ],
            },
        ],
    )

    return response.choices[0].message.content


# ============================================================
# FACT EXTRACTION
# ============================================================

def extract_facts_from_source(
    source,
):
    system_prompt = """
You are a legal case-intelligence extraction agent.

Your task is to extract facts from the supplied source.

IMPORTANT RULES:

1. Extract only information actually supported by the source.
2. Never invent dates, names, locations, events, or evidence.
3. Do not decide whether the applicant is credible.
4. Preserve uncertainty.
5. Distinguish allegations from verified evidence.
6. Every fact must reference the supplied source_id.
7. If a date is uncertain, return null.
8. Do not provide a final legal conclusion.

Return JSON with exactly this structure:

{
  "facts": [
    {
      "statement": "...",
      "date": "YYYY-MM-DD or null",
      "confidence": "low|medium|high"
    }
  ]
}
"""

    user_prompt = f"""
SOURCE ID:
{source["source_id"]}

SOURCE TYPE:
{source["source_type"]}

SOURCE NAME:
{source["name"]}

SOURCE CONTENT:
{source["text"]}
"""

    result = ask_json(
        system_prompt,
        user_prompt,
    )

    created = []

    for item in result.get("facts", []):
        fact = add_fact(
            statement=item.get(
                "statement",
                "",
            ),
            date_value=item.get(
                "date"
            ),
            confidence=item.get(
                "confidence",
                "medium",
            ),
            source_ids=[
                source["source_id"]
            ],
        )

        created.append(fact)

    audit_event(
        "FACT_EXTRACTION",
        details={
            "source_id": source["source_id"],
            "facts_created": len(created),
        },
    )

    return created


# ============================================================
# TIMELINE BUILDER
# ============================================================

def rebuild_timeline():
    case = get_case()

    case["timeline"] = []

    for fact in case["facts"]:

        if not fact.get("date"):
            continue

        add_timeline_event(
            date_value=fact["date"],
            description=fact["statement"],
            fact_ids=[
                fact["fact_id"]
            ],
            source_ids=fact["source_ids"],
            confidence=fact.get(
                "confidence",
                "medium",
            ),
        )

    case["timeline"].sort(
        key=lambda item: (
            item.get("date") or
            "9999-99-99"
        )
    )

    audit_event(
        "TIMELINE_REBUILT",
        details={
            "event_count": len(
                case["timeline"]
            )
        },
    )


# ============================================================
# CONTRADICTION ENGINE
# ============================================================

def run_contradiction_audit():
    case = get_case()

    payload = {
        "facts": case["facts"],
        "timeline": case["timeline"],
        "sources": [
            {
                "source_id": s["source_id"],
                "source_type": s["source_type"],
                "name": s["name"],
            }
            for s in case["sources"]
        ],
    }

    system_prompt = """
You are an adversarial immigration case-quality agent.

Your job is to identify POTENTIAL inconsistencies.

Do NOT accuse an applicant of lying.

Look for:

- conflicting dates
- conflicting locations
- conflicting sequences of events
- different descriptions of the same event
- discrepancies between source documents
- unsupported factual assertions
- chronology problems

For each issue return:

{
  "contradictions": [
    {
      "issue": "...",
      "severity": "low|medium|high|critical",
      "explanation": "...",
      "related_fact_ids": [],
      "recommended_action": "..."
    }
  ]
}

Use "potential inconsistency" unless the conflict is objectively established.
"""

    result = ask_json(
        system_prompt,
        json.dumps(
            payload,
            indent=2,
        ),
    )

    case["vulnerabilities"] = []

    for item in result.get(
        "contradictions",
        [],
    ):
        add_vulnerability(
            issue=item.get(
                "issue",
                "Potential inconsistency",
            ),
            severity=item.get(
                "severity",
                "medium",
            ),
            explanation=item.get(
                "explanation",
                "",
            ),
            fact_ids=item.get(
                "related_fact_ids",
                [],
            ),
            recommended_action=item.get(
                "recommended_action"
            ),
        )

    audit_event(
        "ADVERSARIAL_CONTRADICTION_REVIEW",
        details={
            "issues_found": len(
                case["vulnerabilities"]
            )
        },
    )

    return case["vulnerabilities"]


# ============================================================
# CASE INTELLIGENCE
# ============================================================

def run_case_intelligence(
    protected_ground,
):
    case = get_case()

    facts = case["facts"]

    system_prompt = """
You are an immigration case-intelligence analyst
assisting a supervising attorney.

Analyze the supplied structured case information.

Provide:

1. Case Theory
2. Potential Nexus Strengths
3. Potential Nexus Vulnerabilities
4. Evidence Gaps
5. Timeline / Deadline Concerns
6. Questions the attorney should investigate
7. Facts that require verification

IMPORTANT:

- Do not determine whether the applicant qualifies.
- Do not make credibility findings.
- Do not invent facts.
- Do not cite legal authorities unless supplied.
- Every factual observation must be traceable to a fact_id.
- Clearly distinguish source-supported facts from analysis.
"""

    user_prompt = f"""
Protected Ground:
{protected_ground}

Case ID:
{case["case_id"]}

Applicant:
{case.get("applicant_name", "")}

Country:
{case.get("country_of_origin", "")}

Facts:
{json.dumps(facts, indent=2)}

Timeline:
{json.dumps(case["timeline"], indent=2)}

Evidence:
{json.dumps(case["evidence"], indent=2)}
"""

    result = ask_markdown(
        system_prompt,
        user_prompt,
    )

    case["outputs"].append(
        {
            "type": "case_intelligence",
            "created_at": datetime.now(
                timezone.utc
            ).isoformat(),
            "content": result,
        }
    )

    audit_event(
        "CASE_INTELLIGENCE_ANALYSIS",
        details={
            "protected_ground":
                protected_ground,
        },
    )

    return result


# ============================================================
# SIDEBAR
# ============================================================

st.sidebar.title("⚖️ Primer Paso AI")

st.sidebar.caption(
    "Immigration Case Intelligence & Workflow Suite"
)

page = st.sidebar.radio(
    "Navigation",
    [
        "🏠 Home / Overview",
        "🧠 Case Intelligence",
        "⚡ Live Workflow Tool",
    ],
)


# ============================================================
# CASE STATUS SIDEBAR
# ============================================================

case = get_case()

st.sidebar.divider()

st.sidebar.markdown(
    "### 📁 Current Case"
)

st.sidebar.code(
    case["case_id"]
)

st.sidebar.metric(
    "Facts",
    len(case["facts"]),
)

st.sidebar.metric(
    "Sources",
    len(case["sources"]),
)

st.sidebar.metric(
    "Timeline Events",
    len(case["timeline"]),
)

if case["vulnerabilities"]:
    st.sidebar.warning(
        f"⚠️ {len(case['vulnerabilities'])} "
        "potential issue(s)"
    )

if st.sidebar.button(
    "🔄 Start New Case"
):
    st.session_state.case = create_case()
    st.session_state.last_output = ""
    st.session_state.workflow_result = None
    st.rerun()


# ============================================================
# HOME
# ============================================================

if page == "🏠 Home / Overview":

    st.title(
        "⚖️ Primer Paso AI"
    )

    st.subheader(
        "Immigration Case Intelligence & Workflow Suite"
    )

    st.markdown(
        """
        **Primer Paso AI** is designed to help immigration
        professionals organize, analyze, verify, and prepare
        humanitarian immigration matters while keeping the
        supervising attorney in control.
        """
    )

    st.divider()

    # Product architecture

    st.markdown(
        "## 🧠 Case-Centered Architecture"
    )

    col1, col2, col3 = st.columns(3)

    with col1:

        st.markdown(
            """
            ### 📥 Intake & Ingestion

            - Multilingual document processing
            - Image OCR
            - Client narratives
            - Email intake
            - WhatsApp / voice-note workflow
            - Source tracking
            """
        )

    with col2:

        st.markdown(
            """
            ### 🧩 Case Intelligence

            - Structured facts
            - Case timeline
            - Nexus analysis
            - Country conditions
            - Evidence gaps
            - Deficiency analysis
            - Contradiction detection
            """
        )

    with col3:

        st.markdown(
            """
            ### 🛡️ Verification & Review

            - Citation verification
            - Adversarial case review
            - Exhibit organization
            - Interview preparation
            - Audit trail
            - Attorney approval
            - Case export
            """
        )

    st.divider()

    st.markdown(
        "## 🔐 Product Guardrails"
    )

    g1, g2, g3 = st.columns(3)

    with g1:
        st.info(
            "**Evidence-Grounded**\n\n"
            "Outputs are designed around source material "
            "rather than unsupported factual assumptions."
        )

    with g2:
        st.info(
            "**Human-in-the-Loop**\n\n"
            "AI analysis does not replace supervising "
            "attorney review or professional judgment."
        )

    with g3:
        st.info(
            "**Traceable Case Memory**\n\n"
            "Facts can be connected to their underlying "
            "sources and downstream analysis."
        )

    st.divider()

    st.markdown(
        "## 📊 Current Case Status"
    )

    m1, m2, m3, m4 = st.columns(4)

    m1.metric(
        "Sources",
        len(case["sources"]),
    )

    m2.metric(
        "Facts",
        len(case["facts"]),
    )

    m3.metric(
        "Evidence Items",
        len(case["evidence"]),
    )

    m4.metric(
        "Potential Issues",
        len(case["vulnerabilities"]),
    )

    st.success(
        "Use the sidebar to enter Case Intelligence "
        "or launch the Live Workflow Tool."
    )


# ============================================================
# CASE INTELLIGENCE
# ============================================================

elif page == "🧠 Case Intelligence":

    st.title(
        "🧠 Case Intelligence Center"
    )

    st.caption(
        "Build a structured case record and stress-test it "
        "before attorney review."
    )

    case = get_case()

    # --------------------------------------------------------
    # CASE INFORMATION
    # --------------------------------------------------------

    st.markdown(
        "### 1. Case Information"
    )

    c1, c2 = st.columns(2)

    with c1:

        applicant_name = st.text_input(
            "Applicant Name",
            value=case.get(
                "applicant_name",
                "",
            ),
        )

    with c2:

        country = st.text_input(
            "Country of Origin",
            value=case.get(
                "country_of_origin",
                "",
            ),
        )

    protected_ground = st.selectbox(
        "Primary Protected Ground",
        [""] + PROTECTED_GROUNDS,
        index=(
            [""] + PROTECTED_GROUNDS
        ).index(
            case.get(
                "protected_ground",
                "",
            )
        )
        if case.get(
            "protected_ground",
            "",
        ) in PROTECTED_GROUNDS
        else 0,
    )

    case["applicant_name"] = applicant_name
    case["country_of_origin"] = country
    case["protected_ground"] = protected_ground

    st.divider()

    # --------------------------------------------------------
    # ADD SOURCE
    # --------------------------------------------------------

    st.markdown(
        "### 2. Add Case Source"
    )

    source_type = st.selectbox(
        "Source Type",
        [
            "client_statement",
            "voice_note",
            "whatsapp",
            "document",
            "email",
            "other",
        ],
    )

    source_name = st.text_input(
        "Source Name",
        placeholder="WhatsApp Voice Note #3",
    )

    source_text = st.text_area(
        "Source Text / Transcript",
        height=200,
        placeholder=(
            "Paste a client statement, "
            "voice-note transcript, email, "
            "document text, etc."
        ),
    )

    if st.button(
        "➕ Add Source & Extract Facts",
        type="primary",
    ):

        if not source_text.strip():

            st.warning(
                "Please provide source content."
            )

        else:

            source = add_source(
                source_type=source_type,
                name=(
                    source_name
                    or "Unnamed Source"
                ),
                text=source_text,
            )

            with st.spinner(
                "Extracting source-grounded facts..."
            ):

                try:

                    facts = extract_facts_from_source(
                        source
                    )

                    rebuild_timeline()

                    st.success(
                        f"Source added. "
                        f"{len(facts)} fact(s) extracted."
                    )

                except Exception as exc:

                    st.error(
                        f"Fact extraction failed: {exc}"
                    )

    st.divider()

    # --------------------------------------------------------
    # CASE FACTS
    # --------------------------------------------------------

    st.markdown(
        "### 3. Case Facts"
    )

    if not case["facts"]:

        st.info(
            "No facts have been added yet."
        )

    else:

        for fact in case["facts"]:

            source_display = ", ".join(
                fact["source_ids"]
            )

            with st.expander(
                fact["statement"]
            ):

                st.write(
                    f"**Fact ID:** `{fact['fact_id']}`"
                )

                st.write(
                    f"**Date:** "
                    f"{fact.get('date') or 'Not established'}"
                )

                st.write(
                    f"**Confidence:** "
                    f"{fact.get('confidence', 'medium')}"
                )

                st.write(
                    f"**Source:** "
                    f"`{source_display}`"
                )

                verified = st.checkbox(
                    "Attorney has verified this fact",
                    value=fact.get(
                        "attorney_verified",
                        False,
                    ),
                    key=(
                        f"verify_{fact['fact_id']}"
                    ),
                )

                fact["attorney_verified"] = verified

    st.divider()

    # --------------------------------------------------------
    # TIMELINE
    # --------------------------------------------------------

    st.markdown(
        "### 4. Case Timeline"
    )

    if case["timeline"]:

        for event in case["timeline"]:

            st.markdown(
                f"""
**{event['date']}**

{event['description']}

Source IDs: `{", ".join(event['source_ids'])}`
"""
            )

    else:

        st.info(
            "No dated events have been identified."
        )

    if st.button(
        "🔄 Rebuild Timeline"
    ):

        rebuild_timeline()

        st.success(
            "Timeline rebuilt from structured facts."
        )

    st.divider()

    # --------------------------------------------------------
    # CASE ANALYSIS
    # --------------------------------------------------------

    st.markdown(
        "### 5. Run Case Intelligence"
    )

    if not case["facts"]:

        st.warning(
            "Add case sources and facts first."
        )

    elif not protected_ground:

        st.warning(
            "Select a primary protected ground."
        )

    else:

        if st.button(
            "🧠 Analyze Case",
            type="primary",
        ):

            with st.spinner(
                "Running case intelligence analysis..."
            ):

                try:

                    result = run_case_intelligence(
                        protected_ground
                    )

                    st.session_state.last_output = (
                        result
                    )

                    st.success(
                        "Case intelligence analysis completed."
                    )

                except Exception as exc:

                    st.error(
                        f"Analysis failed: {exc}"
                    )

    if st.session_state.last_output:

        st.markdown(
            "### 📊 Case Intelligence Report"
        )

        st.markdown(
            st.session_state.last_output
        )

    st.divider()

    # --------------------------------------------------------
    # ADVERSARIAL REVIEW
    # --------------------------------------------------------

    st.markdown(
        "### 6. 🔴 Attack This Case"
    )

    st.caption(
        "The purpose of this review is to surface potential "
        "weaknesses for attorney investigation—not to make "
        "credibility findings."
    )

    if st.button(
        "🔴 Run Adversarial Review",
        type="secondary",
    ):

        if not case["facts"]:

            st.warning(
                "There are no facts to review yet."
            )

        else:

            with st.spinner(
                "Running contradiction and vulnerability analysis..."
            ):

                try:

                    issues = run_contradiction_audit()

                    if issues:

                        st.warning(
                            f"{len(issues)} potential "
                            "issue(s) identified."
                        )

                    else:

                        st.success(
                            "No potential contradictions "
                            "were identified by the review."
                        )

                except Exception as exc:

                    st.error(
                        f"Adversarial review failed: {exc}"
                    )

    if case["vulnerabilities"]:

        st.markdown(
            "### 🚨 Potential Vulnerabilities"
        )

        for vulnerability in case[
            "vulnerabilities"
        ]:

            severity = vulnerability[
                "severity"
            ].upper()

            message = (
                f"**{severity} — "
                f"{vulnerability['issue']}**\n\n"
                f"{vulnerability['explanation']}\n\n"
                f"**Recommended action:** "
                f"{vulnerability.get('recommended_action') or 'Attorney review required.'}"
            )

            if severity in [
                "CRITICAL",
                "HIGH",
            ]:

                st.error(message)

            elif severity == "MEDIUM":

                st.warning(message)

            else:

                st.info(message)

    st.divider()

    # --------------------------------------------------------
    # ATTORNEY APPROVAL
    # --------------------------------------------------------

    st.markdown(
        "### 7. 🔒 Attorney Review"
    )

    approval_reason = st.text_area(
        "Attorney review notes",
        placeholder=(
            "Record what was reviewed, "
            "what requires follow-up, "
            "or why the analysis was approved."
        ),
    )

    if st.button(
        "✅ Record Attorney Review",
        type="primary",
    ):

        approval = {
            "approval_id": str(
                uuid.uuid4()
            ),
            "timestamp": datetime.now(
                timezone.utc
            ).isoformat(),
            "actor": "HUMAN_ATTORNEY",
            "notes": approval_reason,
        }

        case["approvals"].append(
            approval
        )

        audit_event(
            "ATTORNEY_REVIEW_RECORDED",
            actor="HUMAN_ATTORNEY",
            details=approval,
        )

        st.success(
            "Attorney review recorded in the case audit trail."
        )


# ============================================================
# LIVE WORKFLOW
# ============================================================

elif page == "⚡ Live Workflow Tool":

    st.title(
        "⚡ Live Workflow Workspace"
    )

    st.caption(
        "Execute specialized immigration workflows while "
        "maintaining a shared case context."
    )

    if client is None:

        st.error(
            "OPENAI_API_KEY is not configured."
        )

        st.info(
            "Add OPENAI_API_KEY to your Streamlit Secrets."
        )

        st.stop()

    workflow_tab = st.selectbox(
        "Select Workflow Module",
        [
            "🌐 Multi-Lingual Translator & Document OCR",
            "⚖️ Nexus & Timeline Auditor",
            "🌍 Country Conditions & Objective Evidence",
            "🔍 Deficiency & Amendment Auditor",
            "📂 Exhibit Index & Document Organizer",
            "✉️ Automated Intake Email Generator",
            "📥 Inbound Firm Mailroom",
        ],
    )

    st.divider()

    # ========================================================
    # TRANSLATOR / OCR
    # ========================================================

    if workflow_tab.startswith(
        "🌐"
    ):

        st.subheader(
            "🌐 Multi-Lingual Translator & Document OCR"
        )

        st.write(
            "Extract text from image documents and "
            "translate source material while preserving "
            "the distinction between original text and analysis."
        )

        uploaded_file = st.file_uploader(
            "Upload image or text document",
            type=[
                "png",
                "jpg",
                "jpeg",
                "txt",
            ],
        )

        target_language = st.selectbox(
            "Target Language",
            TARGET_LANGUAGES,
        )

        if uploaded_file:

            st.success(
                f"Uploaded: {uploaded_file.name}"
            )

            if st.button(
                "🚀 Extract & Translate",
                type="primary",
            ):

                try:

                    file_bytes = (
                        uploaded_file.read()
                    )

                    if uploaded_file.type == (
                        "text/plain"
                    ):

                        source_text = (
                            file_bytes
                            .decode(
                                "utf-8",
                                errors="replace",
                            )
                        )

                        translation_prompt = f"""
Translate the following source text into
{target_language}.

Return:

1. SOURCE TEXT
2. PROFESSIONAL TRANSLATION
3. TRANSLATION NOTES

Do not invent missing text.
Do not provide legal conclusions.

SOURCE:

{source_text}
"""

                        result = ask_markdown(
                            "You are a source-faithful "
                            "legal translation assistant.",
                            translation_prompt,
                        )

                    else:

                        result = process_image(
                            file_bytes,
                            uploaded_file.type,
                            target_language,
                        )

                        source_text = (
                            "[Image OCR source]"
                        )

                    source = add_source(
                        source_type="document",
                        name=uploaded_file.name,
                        text=source_text,
                        metadata={
                            "target_language":
                                target_language,
                            "mime_type":
                                uploaded_file.type,
                        },
                    )

                    st.session_state.last_output = (
                        result
                    )

                    st.success(
                        "OCR / translation completed."
                    )

                    st.markdown(
                        "### 📄 Output"
                    )

                    st.markdown(result)

                    doc = create_docx(
                        f"OCR & Translation — "
                        f"{uploaded_file.name}",
                        result,
                    )

                    st.download_button(
                        "📥 Download DOCX",
                        data=doc,
                        file_name=(
                            f"Translation_"
                            f"{uploaded_file.name}"
                            f".docx"
                        ),
                        mime=(
                            "application/vnd.openxmlformats-"
                            "officedocument.wordprocessingml.document"
                        ),
                    )

                    st.divider()

                    attorney_verified = st.checkbox(
                        "Attorney/staff verified source extraction and translation",
                        key="translation_approval",
                    )

                    if attorney_verified:

                        audit_event(
                            "TRANSLATION_APPROVED",
                            actor="HUMAN_ATTORNEY",
                            details={
                                "source_id":
                                    source[
                                        "source_id"
                                    ]
                            },
                        )

                except Exception as exc:

                    st.error(
                        f"Processing error: {exc}"
                    )

    # ========================================================
    # NEXUS
    # ========================================================

    elif workflow_tab.startswith(
        "⚖️"
    ):

        st.subheader(
            "⚖️ Nexus & Timeline Auditor"
        )

        narrative = st.text_area(
            "Client Narrative",
            height=220,
        )

        ground = st.selectbox(
            "Protected Ground",
            PROTECTED_GROUNDS,
        )

        if st.button(
            "🚀 Run Nexus Audit",
            type="primary",
        ):

            if not narrative.strip():

                st.warning(
                    "Please provide a client narrative."
                )

            else:

                source = add_source(
                    source_type="client_statement",
                    name="Nexus Audit Narrative",
                    text=narrative,
                )

                with st.spinner(
                    "Analyzing narrative..."
                ):

                    try:

                        facts = extract_facts_from_source(
                            source
                        )

                        rebuild_timeline()

                        prompt = f"""
Protected Ground:
{ground}

Facts:
{json.dumps(
    facts,
    indent=2,
)}

Timeline:
{json.dumps(
    get_case()["timeline"],
    indent=2,
)}
"""

                        result = ask_markdown(
                            """
You are an immigration case analyst
assisting a supervising attorney.

Analyze potential nexus strengths,
potential vulnerabilities, evidence gaps,
and timeline concerns.

Do not determine eligibility.
Do not make credibility findings.
Do not invent facts.

Reference fact IDs wherever possible.
""",
                            prompt,
                        )

                        st.markdown(
                            "### 📊 Nexus Analysis"
                        )

                        st.markdown(
                            result
                        )

                        st.session_state.last_output = (
                            result
                        )

                        audit_event(
                            "NEXUS_ANALYSIS",
                            details={
                                "protected_ground":
                                    ground
                            },
                        )

                    except Exception as exc:

                        st.error(
                            f"Nexus analysis failed: {exc}"
                        )

    # ========================================================
    # COUNTRY CONDITIONS
    # ========================================================

    elif workflow_tab.startswith(
        "🌍"
    ):

        st.subheader(
            "🌍 Country Conditions & Objective Evidence"
        )

        home_country = st.text_input(
            "Home Country"
        )

        persecution_category = st.text_input(
            "Primary Persecution / Threat Category"
        )

        key_facts = st.text_area(
            "Key Facts to Corroborate",
            height=160,
        )

        if st.button(
            "🚀 Analyze Country Conditions",
            type="primary",
        ):

            if not all(
                [
                    home_country.strip(),
                    persecution_category.strip(),
                    key_facts.strip(),
                ]
            ):

                st.warning(
                    "Please complete all fields."
                )

            else:

                prompt = f"""
Country:
{home_country}

Threat / Persecution Category:
{persecution_category}

Client Facts:
{key_facts}
"""

                with st.spinner(
                    "Analyzing country-condition issues..."
                ):

                    try:

                        result = ask_markdown(
                            """
You are a country-conditions research
assistant for an immigration attorney.

Identify:

1. Relevant country-condition themes
2. State protection issues
3. Potential corroborating evidence
4. Evidence gaps
5. Questions requiring primary-source research

Do not fabricate sources.
Do not present unsupported current facts
as verified.

The output is a research plan, not a
final legal conclusion.
""",
                            prompt,
                        )

                        st.markdown(
                            "### 🌍 Country Conditions Analysis"
                        )

                        st.markdown(
                            result
                        )

                        audit_event(
                            "COUNTRY_CONDITIONS_ANALYSIS",
                            details={
                                "country":
                                    home_country
                            },
                        )

                    except Exception as exc:

                        st.error(
                            f"Analysis failed: {exc}"
                        )

    # ========================================================
    # DEFICIENCY
    # ========================================================

    elif workflow_tab.startswith(
        "🔍"
    ):

        st.subheader(
            "🔍 Deficiency & Amendment Auditor"
        )

        prior_filing = st.text_area(
            "Original Filing / Declaration",
            height=200,
        )

        government_notice = st.text_area(
            "Government Notice / RFE / Rejection",
            height=150,
        )

        if st.button(
            "🚀 Run Deficiency Analysis",
            type="primary",
        ):

            if not prior_filing.strip():

                st.warning(
                    "Please provide the prior filing."
                )

            else:

                with st.spinner(
                    "Reviewing filing for potential deficiencies..."
                ):

                    try:

                        result = ask_markdown(
                            """
You are an immigration filing-review
assistant.

Identify:

1. Factual gaps
2. Internal inconsistencies
3. Missing corroboration
4. Potential legal issues requiring attorney research
5. Questions raised by the government notice
6. Recommended follow-up investigation

Do not invent facts.
Do not guarantee an outcome.
Do not make a final legal determination.
""",
                            f"""
PRIOR FILING:

{prior_filing}

GOVERNMENT NOTICE:

{government_notice}
""",
                        )

                        st.markdown(
                            "### 🔍 Deficiency Report"
                        )

                        st.markdown(
                            result
                        )

                        audit_event(
                            "DEFICIENCY_REVIEW",
                        )

                    except Exception as exc:

                        st.error(
                            f"Deficiency analysis failed: {exc}"
                        )

    # ========================================================
    # EXHIBITS
    # ========================================================

    elif workflow_tab.startswith(
        "📂"
    ):

        st.subheader(
            "📂 Exhibit Index & Document Organizer"
        )

        documents = st.text_area(
            "Case Documents",
            height=220,
            placeholder=(
                "1. Birth certificate\n"
                "2. Police report\n"
                "3. Medical records\n"
                "4. Country-condition report"
            ),
        )

        if st.button(
            "🚀 Generate Exhibit Index",
            type="primary",
        ):

            if not documents.strip():

                st.warning(
                    "Please list the case documents."
                )

            else:

                with st.spinner(
                    "Organizing exhibits..."
                ):

                    try:

                        result = ask_markdown(
                            """
You are an immigration legal
document-organizing assistant.

Create a clean exhibit index.

Do not invent dates.
Do not invent document contents.
If information is missing, label it
"Not provided."

Suggested columns:

Exhibit
Document Description
Date
Evidentiary Purpose
Verification Status
""",
                            documents,
                        )

                        st.markdown(
                            "### 📋 Exhibit Index"
                        )

                        st.markdown(
                            result
                        )

                        doc = create_docx(
                            "Master Exhibit Index",
                            result,
                        )

                        st.download_button(
                            "📥 Download Exhibit Index",
                            data=doc,
                            file_name=(
                                "Master_Exhibit_Index.docx"
                            ),
                            mime=(
                                "application/vnd.openxmlformats-"
                                "officedocument.wordprocessingml.document"
                            ),
                        )

                        audit_event(
                            "EXHIBIT_INDEX_GENERATED"
                        )

                    except Exception as exc:

                        st.error(
                            f"Exhibit generation failed: {exc}"
                        )

    # ========================================================
    # EMAIL
    # ========================================================

    elif workflow_tab.startswith(
        "✉️"
    ):

        st.subheader(
            "✉️ Automated Intake Email Generator"
        )

        client_name = st.text_input(
            "Client Name"
        )

        language = st.selectbox(
            "Preferred Language",
            TARGET_LANGUAGES,
        )

        purpose = st.selectbox(
            "Email Purpose",
            [
                "Document Request",
                "Additional Evidence Request",
                "Interview Reminder",
                "Case Status Update",
                "Custom",
            ],
        )

        details = st.text_area(
            "Specific Details",
            height=150,
        )

        if purpose == "Custom":

            purpose = st.text_input(
                "Custom Purpose"
            )

        if st.button(
            "🚀 Generate Email",
            type="primary",
        ):

            if not all(
                [
                    client_name.strip(),
                    purpose.strip(),
                    details.strip(),
                ]
            ):

                st.warning(
                    "Please complete all fields."
                )

            else:

                with st.spinner(
                    "Drafting client correspondence..."
                ):

                    try:

                        result = ask_markdown(
                            f"""
You are an immigration law-office
communications assistant.

Write a professional, clear,
empathetic client email in
{language}.

Do not provide legal advice beyond
the supplied information.

Clearly identify requested actions.
Return only the email draft.
""",
                            f"""
Client:
{client_name}

Purpose:
{purpose}

Details:
{details}
""",
                        )

                        st.markdown(
                            "### 📨 Draft Email"
                        )

                        st.markdown(
                            result
                        )

                        doc = create_docx(
                            f"Client Correspondence — {client_name}",
                            result,
                        )

                        st.download_button(
                            "📥 Download Email",
                            data=doc,
                            file_name=(
                                "Client_Correspondence.docx"
                            ),
                            mime=(
                                "application/vnd.openxmlformats-"
                                "officedocument.wordprocessingml.document"
                            ),
                        )

                        approved = st.checkbox(
                            "Staff/attorney reviewed this draft before sending",
                            key="email_approval",
                        )

                        if approved:

                            audit_event(
                                "CLIENT_EMAIL_APPROVED",
                                actor="HUMAN_ATTORNEY",
                            )

                    except Exception as exc:

                        st.error(
                            f"Email generation failed: {exc}"
                        )

    # ========================================================
    # MAILROOM
    # ========================================================

    elif workflow_tab.startswith(
        "📥"
    ):

        st.subheader(
            "📥 Inbound Firm Mailroom"
        )

        st.info(
            "This simulation lets you test inbound "
            "client-message classification before "
            "connecting a production inbox."
        )

        sender = st.text_input(
            "Sender",
            value="client@example.com",
        )

        subject = st.text_input(
            "Subject",
            value="Question about my case",
        )

        message = st.text_area(
            "Message",
            height=180,
        )

        if st.button(
            "🚀 Process Message",
            type="primary",
        ):

            if not message.strip():

                st.warning(
                    "Please provide a message."
                )

            else:

                with st.spinner(
                    "Classifying inbound message..."
                ):

                    try:

                        result = ask_markdown(
                            """
You are an immigration law-office
mailroom assistant.

Analyze the incoming message.

Return:

1. Client intent
2. Urgency
3. Potential deadline concern
4. Documents requested or missing
5. Questions requiring staff review
6. Draft response

Do not make legal conclusions.
Do not promise outcomes.
""",
                            f"""
FROM:
{sender}

SUBJECT:
{subject}

MESSAGE:
{message}
""",
                        )

                        st.markdown(
                            "### 📨 Mailroom Analysis"
                        )

                        st.markdown(
                            result
                        )

                        audit_event(
                            "INBOUND_MESSAGE_PROCESSED",
                            details={
                                "sender":
                                    sender,
                                "subject":
                                    subject,
                            },
                        )

                    except Exception as exc:

                        st.error(
                            f"Mailroom processing failed: {exc}"
                        )


# ============================================================
# FOOTER
# ============================================================

st.divider()

st.caption(
    "Primer Paso AI — AI-assisted immigration workflow. "
    "Outputs require appropriate professional review."
)
  
                        
                          
